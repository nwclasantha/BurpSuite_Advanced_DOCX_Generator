#!/usr/bin/env python3
import argparse
import logging
import os
import re
import sys
from collections import defaultdict, Counter
from dataclasses import dataclass, field
from datetime import datetime
from typing import List, Optional, Tuple, Dict
from urllib.parse import urlparse

try:
    from bs4 import BeautifulSoup, Tag
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml, OxmlElement
    from docx.table import Table
    import pandas as pd
except ImportError:
    print("═" * 60)
    print("  ERROR: Required packages not installed.")
    print("═" * 60)
    print("  Run: pip install python-docx beautifulsoup4 lxml pandas openpyxl")
    print("═" * 60)
    sys.exit(1)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class VulnerabilityInstance:
    """Single vulnerability instance."""
    url: str
    host: str = ""
    path: str = ""
    severity: str = "Information"
    confidence: str = "Tentative"
    issue_detail: Optional[str] = None
    request: Optional[str] = None
    response: Optional[str] = None
    request_highlights: List[str] = field(default_factory=list)
    response_highlights: List[str] = field(default_factory=list)


@dataclass
class Vulnerability:
    """Vulnerability type with instances."""
    title: str
    severity: str = "Information"
    confidence: str = "Tentative"
    issue_background: Optional[str] = None
    issue_remediation: Optional[str] = None
    references: List[Tuple[str, str]] = field(default_factory=list)
    cwe_classifications: List[Tuple[str, str]] = field(default_factory=list)
    instances: List[VulnerabilityInstance] = field(default_factory=list)
    section_id: str = ""


@dataclass
class ReportStatistics:
    """Report statistics."""
    high_certain: int = 0
    high_firm: int = 0
    high_tentative: int = 0
    medium_certain: int = 0
    medium_firm: int = 0
    medium_tentative: int = 0
    low_certain: int = 0
    low_firm: int = 0
    low_tentative: int = 0
    info_certain: int = 0
    info_firm: int = 0
    info_tentative: int = 0

    @property
    def total_high(self) -> int:
        return self.high_certain + self.high_firm + self.high_tentative

    @property
    def total_medium(self) -> int:
        return self.medium_certain + self.medium_firm + self.medium_tentative

    @property
    def total_low(self) -> int:
        return self.low_certain + self.low_firm + self.low_tentative

    @property
    def total_info(self) -> int:
        return self.info_certain + self.info_firm + self.info_tentative

    @property
    def total_issues(self) -> int:
        return self.total_high + self.total_medium + self.total_low + self.total_info


# ============================================================================
# COLORS AND STYLING
# ============================================================================

class Colors:
    """Premium color palette for enterprise reports."""
    # Severity Colors - Professional gradient-inspired
    CRITICAL = RGBColor(139, 0, 0)      # Dark red
    CRITICAL_BG = RGBColor(255, 230, 230)
    HIGH = RGBColor(192, 0, 0)          # Strong red
    HIGH_BG = RGBColor(255, 235, 235)
    MEDIUM = RGBColor(255, 140, 0)      # Dark orange
    MEDIUM_BG = RGBColor(255, 248, 235)
    LOW = RGBColor(0, 112, 192)         # Professional blue
    LOW_BG = RGBColor(235, 245, 255)
    INFO = RGBColor(89, 89, 89)         # Charcoal gray
    INFO_BG = RGBColor(245, 245, 245)

    # Enterprise Document Colors
    PRIMARY = RGBColor(0, 51, 102)       # Navy blue
    PRIMARY_LIGHT = RGBColor(0, 82, 147) # Lighter navy
    SECONDARY = RGBColor(51, 51, 51)     # Dark charcoal
    ACCENT = RGBColor(0, 120, 215)       # Accent blue
    ACCENT_DARK = RGBColor(0, 90, 158)   # Darker accent

    # Neutral Colors
    WHITE = RGBColor(255, 255, 255)
    BLACK = RGBColor(0, 0, 0)
    GRAY = RGBColor(128, 128, 128)
    LIGHT_GRAY = RGBColor(245, 245, 245)
    MEDIUM_GRAY = RGBColor(200, 200, 200)
    DARK_GRAY = RGBColor(64, 64, 64)

    # Status Colors
    SUCCESS = RGBColor(34, 139, 34)      # Forest green
    SUCCESS_BG = RGBColor(235, 255, 235)
    WARNING = RGBColor(255, 165, 0)      # Orange
    WARNING_BG = RGBColor(255, 250, 235)

    # Special Elements
    CODE_BG = RGBColor(248, 249, 250)
    CODE_BORDER = RGBColor(225, 228, 232)
    BORDER = RGBColor(206, 212, 218)
    TABLE_HEADER = RGBColor(0, 51, 102)
    TABLE_ALT_ROW = RGBColor(250, 250, 252)

    # Risk Score Colors
    RISK_CRITICAL = RGBColor(139, 0, 0)
    RISK_HIGH = RGBColor(192, 0, 0)
    RISK_MEDIUM = RGBColor(255, 140, 0)
    RISK_LOW = RGBColor(0, 128, 0)


class Fonts:
    """Professional font stack."""
    TITLE = "Segoe UI"
    HEADING = "Segoe UI"
    BODY = "Calibri"
    CODE = "Consolas"
    ACCENT = "Calibri Light"


# OWASP Top 10 2021 Mapping
OWASP_MAPPING = {
    'injection': ('A03:2021', 'Injection'),
    'sql injection': ('A03:2021', 'Injection'),
    'command injection': ('A03:2021', 'Injection'),
    'xss': ('A03:2021', 'Injection'),
    'cross-site scripting': ('A03:2021', 'Injection'),
    'broken access': ('A01:2021', 'Broken Access Control'),
    'authorization': ('A01:2021', 'Broken Access Control'),
    'idor': ('A01:2021', 'Broken Access Control'),
    'cryptographic': ('A02:2021', 'Cryptographic Failures'),
    'ssl': ('A02:2021', 'Cryptographic Failures'),
    'tls': ('A02:2021', 'Cryptographic Failures'),
    'certificate': ('A02:2021', 'Cryptographic Failures'),
    'encryption': ('A02:2021', 'Cryptographic Failures'),
    'security misconfiguration': ('A05:2021', 'Security Misconfiguration'),
    'header': ('A05:2021', 'Security Misconfiguration'),
    'cors': ('A05:2021', 'Security Misconfiguration'),
    'vulnerable component': ('A06:2021', 'Vulnerable Components'),
    'outdated': ('A06:2021', 'Vulnerable Components'),
    'authentication': ('A07:2021', 'Auth Failures'),
    'session': ('A07:2021', 'Auth Failures'),
    'password': ('A07:2021', 'Auth Failures'),
    'csrf': ('A01:2021', 'Broken Access Control'),
    'ssrf': ('A10:2021', 'SSRF'),
    'xxe': ('A05:2021', 'Security Misconfiguration'),
    'deserialization': ('A08:2021', 'Software/Data Integrity'),
    'logging': ('A09:2021', 'Logging Failures'),
    'monitoring': ('A09:2021', 'Logging Failures'),
}


def get_owasp_category(vuln_title: str) -> Tuple[str, str]:
    """Map vulnerability to OWASP Top 10 category."""
    title_lower = vuln_title.lower()
    for keyword, mapping in OWASP_MAPPING.items():
        if keyword in title_lower:
            return mapping
    return ('A05:2021', 'Security Misconfiguration')  # Default


# ============================================================================
# HTML PARSER
# ============================================================================

class BurpHTMLParser:
    """Parses BurpSuite HTML reports."""

    def __init__(self, html_content: str):
        logger.info("Initializing HTML parser...")
        self.soup = BeautifulSoup(html_content, 'lxml')
        self.statistics = ReportStatistics()
        self.vulnerabilities: List[Vulnerability] = []

    def parse(self) -> Tuple[ReportStatistics, List[Vulnerability]]:
        """Parse the complete HTML report."""
        logger.info("Parsing HTML report...")
        self._extract_statistics()
        self._extract_vulnerabilities()

        total_instances = sum(len(v.instances) for v in self.vulnerabilities)
        logger.info(f"Parsed {len(self.vulnerabilities)} vulnerability types, {total_instances} instances")
        return self.statistics, self.vulnerabilities

    def _extract_statistics(self):
        """Extract statistics from summary table."""
        overview_table = self.soup.find('table', class_='overview_table')
        if not overview_table:
            return

        mapping = {
            'high_certain': 'high_certain', 'high_firm': 'high_firm', 'high_tentative': 'high_tentative',
            'medium_certain': 'medium_certain', 'medium_firm': 'medium_firm', 'medium_tentative': 'medium_tentative',
            'low_certain': 'low_certain', 'low_firm': 'low_firm', 'low_tentative': 'low_tentative',
            'info_certain': 'info_certain', 'info_firm': 'info_firm', 'info_tentative': 'info_tentative',
        }

        for class_name, attr_name in mapping.items():
            elem = overview_table.find('span', class_=f'colour_block {class_name}')
            if elem:
                try:
                    setattr(self.statistics, attr_name, int(elem.get_text(strip=True)))
                except ValueError:
                    pass

    def _extract_vulnerabilities(self):
        """Extract all vulnerabilities."""
        bodh0_elements = self.soup.find_all('span', class_='BODH0')
        total = len(bodh0_elements)

        for idx, bodh0 in enumerate(bodh0_elements):
            if idx % 10 == 0:
                logger.info(f"Processing vulnerability {idx + 1}/{total}...")
            vuln = self._parse_vulnerability(bodh0)
            if vuln:
                self.vulnerabilities.append(vuln)

    def _parse_vulnerability(self, bodh0_elem) -> Optional[Vulnerability]:
        """Parse single vulnerability section."""
        title_text = bodh0_elem.get_text(strip=True)
        title_match = re.match(r'^\d+\.\s*(.+)$', title_text)
        title = title_match.group(1) if title_match else title_text

        vuln = Vulnerability(title=title, section_id=bodh0_elem.get('id', ''))

        # Navigate through content - find the outermost container
        current = bodh0_elem
        while current.parent and isinstance(current.parent, Tag) and current.parent.name != 'body':
            if current.parent.name in ['div', 'span'] and current.find_previous_sibling() is None:
                current = current.parent
            else:
                break

        content_elements = []
        sibling = current.next_sibling
        while sibling:
            if isinstance(sibling, Tag):
                if sibling.find('span', class_='BODH0') or (sibling.name == 'span' and 'BODH0' in sibling.get('class', [])):
                    break
                content_elements.append(sibling)
            sibling = sibling.next_sibling

        self._parse_content(vuln, content_elements)
        return vuln

    def _parse_content(self, vuln: Vulnerability, elements: list):
        """Parse vulnerability content."""
        current_section = None
        current_instance = None
        instances = []

        for elem in elements:
            if not isinstance(elem, Tag):
                continue

            # Check for instance header (BODH1)
            bodh1 = elem.find('span', class_='BODH1') if hasattr(elem, 'find') else None
            if not bodh1 and elem.name == 'span' and 'BODH1' in elem.get('class', []):
                bodh1 = elem

            if bodh1:
                if current_instance:
                    instances.append(current_instance)
                instance_text = bodh1.get_text(strip=True)
                url_match = re.match(r'^\d+\.\d+\.\s*(.+)$', instance_text)
                url = url_match.group(1) if url_match else instance_text
                current_instance = VulnerabilityInstance(url=url)
                current_section = None
                continue

            # Check for section headers (h2)
            h2 = elem.find('h2') if hasattr(elem, 'find') else None
            if not h2 and elem.name == 'h2':
                h2 = elem
            if h2:
                current_section = h2.get_text(strip=True).lower()
                continue

            # Extract summary table metadata - check if elem IS the table or contains it
            summary_table = None
            if elem.name == 'table' and 'summary_table' in elem.get('class', []):
                summary_table = elem
            elif hasattr(elem, 'find'):
                summary_table = elem.find('table', class_='summary_table')

            if summary_table and current_instance:
                self._extract_metadata(summary_table, current_instance)
                continue

            # Extract TEXT content
            text_span = elem.find('span', class_='TEXT') if hasattr(elem, 'find') else None
            if not text_span and hasattr(elem, 'get') and 'TEXT' in elem.get('class', []):
                text_span = elem

            if text_span:
                content = self._get_text(text_span)
                if current_section == 'issue background':
                    vuln.issue_background = content
                elif current_section == 'issue remediation':
                    vuln.issue_remediation = content
                elif current_section == 'references':
                    if hasattr(text_span, 'find_all'):
                        vuln.references = [(a.get_text(strip=True), a.get('href', '')) for a in text_span.find_all('a')]
                elif current_section == 'vulnerability classifications':
                    if hasattr(text_span, 'find_all'):
                        vuln.cwe_classifications = [(a.get_text(strip=True), a.get('href', '')) for a in text_span.find_all('a')]
                elif current_instance and current_section == 'issue detail':
                    current_instance.issue_detail = content

            # Extract request/response
            # Use recursive=False to avoid finding nested rr_divs multiple times
            rr_divs = elem.find_all('div', class_='rr_div', recursive=False) if hasattr(elem, 'find_all') else []
            if not rr_divs and elem.name == 'div' and 'rr_div' in elem.get('class', []):
                rr_divs = [elem]
            # If no direct children found, try recursive search as fallback (for deeply nested structures)
            if not rr_divs and hasattr(elem, 'find_all'):
                rr_divs = elem.find_all('div', class_='rr_div')

            for rr_div in rr_divs:
                if current_instance:
                    content, highlights = self._extract_rr(rr_div)
                    # BurpSuite uses "Request 1", "Response 1" etc. as h2 headers
                    if current_section and current_section.startswith('request'):
                        # Append content if multiple rr_divs for same section
                        if current_instance.request:
                            current_instance.request += '\n' + content
                        else:
                            current_instance.request = content
                        # Extend highlights instead of overwriting to collect all unique evidence
                        for h in highlights:
                            if h not in current_instance.request_highlights:
                                current_instance.request_highlights.append(h)
                    elif current_section and current_section.startswith('response'):
                        if current_instance.response:
                            current_instance.response += '\n' + content
                        else:
                            current_instance.response = content
                        for h in highlights:
                            if h not in current_instance.response_highlights:
                                current_instance.response_highlights.append(h)

        if current_instance:
            instances.append(current_instance)

        vuln.instances = instances
        if instances:
            vuln.severity = instances[0].severity
            vuln.confidence = instances[0].confidence

    def _get_text(self, elem) -> str:
        """Extract text from element without modifying the original."""
        # Create a copy to avoid modifying the original element
        elem_copy = BeautifulSoup(str(elem), 'lxml')
        for br in elem_copy.find_all('br'):
            br.replace_with('\n')
        lines = [line.strip() for line in elem_copy.get_text().split('\n')]
        return '\n'.join(line for line in lines if line)

    def _extract_metadata(self, table, instance: VulnerabilityInstance):
        """Extract instance metadata."""
        rows = table.find_all('tr')
        for row in rows:
            cells = row.find_all('td')
            # Iterate through cells looking for label:value pairs
            i = 0
            while i < len(cells):
                cell_text = cells[i].get_text(strip=True).lower().replace('\xa0', ' ')

                if 'severity' in cell_text:
                    # Next cell contains the value
                    if i + 1 < len(cells):
                        instance.severity = cells[i + 1].get_text(strip=True)
                        i += 2
                        continue
                elif 'confidence' in cell_text:
                    if i + 1 < len(cells):
                        instance.confidence = cells[i + 1].get_text(strip=True)
                        i += 2
                        continue
                elif 'host' in cell_text:
                    if i + 1 < len(cells):
                        instance.host = cells[i + 1].get_text(strip=True)
                        i += 2
                        continue
                elif 'path' in cell_text:
                    if i + 1 < len(cells):
                        instance.path = cells[i + 1].get_text(strip=True)
                        i += 2
                        continue
                i += 1

    def _extract_rr(self, div) -> Tuple[str, List[str]]:
        """Extract request/response content."""
        # Get all highlights, strip whitespace, and deduplicate while preserving order
        all_highlights = [h.get_text().strip() for h in div.find_all('span', class_='HIGHLIGHT')]
        seen = set()
        highlights = []
        for h in all_highlights:
            # Skip empty highlights and duplicates
            if h and h not in seen:
                seen.add(h)
                highlights.append(h)
        div_copy = BeautifulSoup(str(div), 'lxml')
        for br in div_copy.find_all('br'):
            br.replace_with('\n')
        return div_copy.get_text().strip(), highlights


# ============================================================================
# DOCX GENERATOR - ENTERPRISE PROFESSIONAL
# ============================================================================

class DocxReportGenerator:
    """Enterprise-grade DOCX report generator."""

    # Severity mapping - BurpSuite uses "High" as max (no "Critical")
    # We map "critical" input to "High" for user convenience
    SEVERITY_MAP = {
        'critical': 'High',
        'high': 'High',
        'medium': 'Medium',
        'low': 'Low',
        'info': 'Information',
        'information': 'Information',
    }

    SEVERITY_ORDER = {'High': 0, 'Medium': 1, 'Low': 2, 'Information': 3}

    def __init__(self, statistics: ReportStatistics, vulnerabilities: List[Vulnerability],
                 company_name: str, report_title: str, target_name: str,
                 severity_filter: List[str] = None, evidence_only: bool = False,
                 network_reports_folder: str = None):
        self.stats = statistics
        self.vulns = vulnerabilities
        self.company_name = company_name
        self.report_title = report_title
        self.target_name = target_name
        self.doc = Document()
        self.evidence_only = evidence_only  # Show only evidence, not full request/response
        self.network_reports_folder = network_reports_folder  # Network pentest Excel files folder
        self.network_vulns = {}  # Parsed network vulnerabilities by domain

        # Process severity filter
        if severity_filter:
            self.severity_filter = set()
            invalid_severities = []
            for sev in severity_filter:
                mapped = self.SEVERITY_MAP.get(sev.lower().strip())
                if mapped:
                    self.severity_filter.add(mapped)
                else:
                    invalid_severities.append(sev)

            # Warn about invalid severity values
            if invalid_severities:
                logger.warning(f"Invalid severity values ignored: {', '.join(invalid_severities)}")
                logger.warning(f"Valid options: critical, high, medium, low, info")

            if not self.severity_filter:
                logger.warning("No valid severity filter provided, using all severities")
                self.severity_filter = {'High', 'Medium', 'Low', 'Information'}
        else:
            # Default: all severities
            self.severity_filter = {'High', 'Medium', 'Low', 'Information'}

        # Filter vulnerabilities based on severity
        self.filtered_vulns = [v for v in self.vulns if v.severity in self.severity_filter]
        self.filtered_vulns = sorted(self.filtered_vulns,
                                      key=lambda v: self.SEVERITY_ORDER.get(v.severity, 4))

    def generate(self, output_path: str):
        """Generate enterprise DOCX report."""
        logger.info("Generating enterprise DOCX report...")
        logger.info(f"Severity filter: {', '.join(sorted(self.severity_filter))}")
        logger.info(f"Filtered vulnerabilities: {len(self.filtered_vulns)} types")

        # Warn if no vulnerabilities match the filter
        if not self.filtered_vulns:
            logger.warning("WARNING: No vulnerabilities match the selected severity filter!")
            logger.warning(f"Filter was: {', '.join(sorted(self.severity_filter))}")

        self._setup_document()
        self._add_header_footer()
        self._create_cover_page()
        self._create_table_of_contents()  # Table of Contents
        self._create_executive_dashboard()
        self._create_visual_charts()  # Visual bar charts
        self._create_owasp_compliance()  # OWASP Top 10 Mapping
        self._create_remediation_priority()  # Remediation Priority Matrix (Section 4)
        self._create_risk_matrix()  # Risk Assessment Matrix (Section 5)
        self._create_findings_summary()
        self._create_detailed_findings()
        self._create_appendix()

        # Add Network Pentest section if folder provided
        if self.network_reports_folder and os.path.isdir(self.network_reports_folder):
            self._parse_network_reports()
            self._create_network_assessment_section()

        logger.info("Saving document...")
        try:
            self.doc.save(output_path)
            logger.info(f"Report saved: {output_path}")
        except PermissionError:
            logger.error(f"Permission denied: Cannot write to {output_path}")
            logger.error("Please close the file if it's open in another application")
            raise
        except OSError as e:
            logger.error(f"Error saving file: {e}")
            raise

    def _setup_document(self):
        """Setup document."""
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)
            section.header_distance = Cm(1)
            section.footer_distance = Cm(1)

    def _add_header_footer(self):
        """Add professional header and footer with Page X of Y numbering."""
        section = self.doc.sections[0]

        # Header with bottom border
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = header_para.add_run(f"{self.company_name} | {self.report_title}")
        run.font.name = Fonts.BODY
        run.font.size = Pt(9)
        run.font.color.rgb = Colors.GRAY
        run.font.italic = True

        # Add bottom border to header
        header_para.paragraph_format.space_after = Pt(6)
        self._add_paragraph_border(header_para, bottom=True)

        # Footer with Page X of Y
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add top border to footer
        self._add_paragraph_border(footer_para, top=True)
        footer_para.paragraph_format.space_before = Pt(6)

        run = footer_para.add_run("CONFIDENTIAL")
        run.font.name = Fonts.BODY
        run.font.size = Pt(8)
        run.font.color.rgb = Colors.HIGH
        run.font.bold = True

        run = footer_para.add_run(" | ")
        run.font.size = Pt(8)

        # Add Page X of Y using Word field codes
        self._add_page_number_field(footer_para)

        run = footer_para.add_run(" | ")
        run.font.size = Pt(8)

        run = footer_para.add_run(f"{datetime.now().strftime('%Y-%m-%d')}")
        run.font.name = Fonts.BODY
        run.font.size = Pt(8)
        run.font.color.rgb = Colors.GRAY

    def _add_page_number_field(self, paragraph):
        """Add Page X of Y field to paragraph."""
        # "Page " text
        run = paragraph.add_run("Page ")
        run.font.name = Fonts.BODY
        run.font.size = Pt(8)
        run.font.color.rgb = Colors.GRAY

        # PAGE field (current page number)
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = Fonts.BODY
        run.font.size = Pt(8)
        run.font.color.rgb = Colors.GRAY

        # " of " text
        run = paragraph.add_run(" of ")
        run.font.name = Fonts.BODY
        run.font.size = Pt(8)
        run.font.color.rgb = Colors.GRAY

        # NUMPAGES field (total pages)
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "NUMPAGES"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = Fonts.BODY
        run.font.size = Pt(8)
        run.font.color.rgb = Colors.GRAY

    def _add_paragraph_border(self, paragraph, top=False, bottom=False):
        """Add border to paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        if bottom:
            bottom_bdr = OxmlElement('w:bottom')
            bottom_bdr.set(qn('w:val'), 'single')
            bottom_bdr.set(qn('w:sz'), '6')
            bottom_bdr.set(qn('w:space'), '1')
            bottom_bdr.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom_bdr)

        if top:
            top_bdr = OxmlElement('w:top')
            top_bdr.set(qn('w:val'), 'single')
            top_bdr.set(qn('w:sz'), '6')
            top_bdr.set(qn('w:space'), '1')
            top_bdr.set(qn('w:color'), 'CCCCCC')
            pBdr.append(top_bdr)

        pPr.append(pBdr)

    def _create_table_of_contents(self):
        """Create Table of Contents page with auto-updating TOC field."""
        # TOC Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TABLE OF CONTENTS")
        run.font.name = Fonts.HEADING
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY
        p.paragraph_format.space_after = Pt(24)

        # Add instruction note
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("(Right-click and select 'Update Field' to refresh page numbers)")
        run.font.name = Fonts.BODY
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = Colors.GRAY
        p.paragraph_format.space_after = Pt(18)

        # Add TOC field code - this creates an auto-updating TOC in Word
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()

        # Begin field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        # TOC instruction - \o "1-3" means include heading levels 1-3, \h creates hyperlinks
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

        # Separate field
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        # End field
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Add placeholder text (will be replaced when TOC is updated in Word)
        placeholder_run = paragraph.add_run("Update this field to generate Table of Contents")
        placeholder_run.font.name = Fonts.BODY
        placeholder_run.font.size = Pt(11)
        placeholder_run.font.color.rgb = Colors.GRAY

        # Close the field
        end_run = paragraph.add_run()
        end_run._r.append(fldChar3)

        # Add manual TOC entries as backup (visible before field update)
        self.doc.add_paragraph()
        self._add_toc_entry("1. EXECUTIVE SUMMARY", 1)
        self._add_toc_entry("   1.1 Assessment Scope", 2)
        self._add_toc_entry("   1.2 Key Findings Overview", 2)
        self._add_toc_entry("   1.3 Immediate Action Required", 2)
        self._add_toc_entry("2. SEVERITY DISTRIBUTION CHARTS", 1)
        self._add_toc_entry("   2.1 Findings by Severity Level", 2)
        self._add_toc_entry("   2.2 Findings by Confidence Level", 2)
        self._add_toc_entry("3. OWASP TOP 10 COMPLIANCE MAPPING", 1)
        self._add_toc_entry("4. REMEDIATION PRIORITY MATRIX", 1)
        self._add_toc_entry("5. RISK ASSESSMENT MATRIX", 1)
        self._add_toc_entry("6. VULNERABILITY SUMMARY", 1)
        self._add_toc_entry("7. DETAILED FINDINGS", 1)

        # Add vulnerability entries to TOC
        for idx, vuln in enumerate(self.filtered_vulns[:10], 1):  # Limit to first 10
            title = vuln.title[:45] + '...' if len(vuln.title) > 45 else vuln.title
            self._add_toc_entry(f"   7.{idx} {title}", 2)

        if len(self.filtered_vulns) > 10:
            self._add_toc_entry(f"   ... and {len(self.filtered_vulns) - 10} more findings", 2)

        self._add_toc_entry("8. APPENDIX", 1)
        self._add_toc_entry("   8.1 Severity Definitions", 2)
        self._add_toc_entry("   8.2 Confidence Levels", 2)
        self._add_toc_entry("   8.3 Assessment Methodology", 2)

        # Add Section 9 if network reports folder is provided
        if self.network_reports_folder and os.path.isdir(self.network_reports_folder):
            self._add_toc_entry("9. NETWORK VULNERABILITY ASSESSMENT", 1)
            self._add_toc_entry("   9.1 Network Vulnerability Summary", 2)
            self._add_toc_entry("   9.2 Detailed Network Findings by Host", 2)

        self.doc.add_page_break()

    def _add_toc_entry(self, text: str, level: int):
        """Add a manual TOC entry."""
        p = self.doc.add_paragraph()

        if level == 1:
            run = p.add_run(text)
            run.font.name = Fonts.BODY
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = Colors.PRIMARY
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        else:
            run = p.add_run(text)
            run.font.name = Fonts.BODY
            run.font.size = Pt(10)
            run.font.color.rgb = Colors.SECONDARY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        # Add dotted leader and page placeholder
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))

    def _create_cover_page(self):
        """Create premium enterprise cover page with visual elements."""
        # ═══════════════════════════════════════════════════════════════════
        # TOP HEADER BAR
        # ═══════════════════════════════════════════════════════════════════
        header_table = self.doc.add_table(rows=1, cols=1)
        header_cell = header_table.rows[0].cells[0]
        self._set_cell_bg(header_cell, Colors.PRIMARY)
        header_table.rows[0].height = Twips(800)

        p = header_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("▰▰▰  CONFIDENTIAL SECURITY ASSESSMENT  ▰▰▰")
        run.font.name = Fonts.TITLE
        run.font.size = Pt(12)
        run.font.color.rgb = Colors.WHITE
        run.font.bold = True

        self.doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════
        # COMPANY BRANDING
        # ═══════════════════════════════════════════════════════════════════
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = Colors.PRIMARY
        run.font.size = Pt(8)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.company_name.upper())
        run.font.name = Fonts.TITLE
        run.font.size = Pt(24)
        run.font.color.rgb = Colors.PRIMARY
        run.font.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = Colors.PRIMARY
        run.font.size = Pt(8)

        self.doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════
        # MAIN TITLE
        # ═══════════════════════════════════════════════════════════════════
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.report_title.upper())
        run.font.name = Fonts.TITLE
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Enterprise Security Assessment Report")
        run.font.name = Fonts.ACCENT
        run.font.size = Pt(14)
        run.font.color.rgb = Colors.SECONDARY

        self.doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════
        # RISK SCORE GAUGE (Visual representation)
        # ═══════════════════════════════════════════════════════════════════
        risk_score, risk_level, risk_color = self._calculate_risk_score()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("◆ ENTERPRISE RISK SCORE ◆")
        run.font.name = Fonts.HEADING
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = Colors.DARK_GRAY

        # Create visual gauge
        gauge_table = self.doc.add_table(rows=2, cols=1)
        gauge_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        gauge_cell = gauge_table.rows[0].cells[0]
        gauge_cell.width = Inches(4)

        # Score display
        p = gauge_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{risk_score}")
        run.font.name = Fonts.TITLE
        run.font.size = Pt(48)
        run.font.bold = True
        run.font.color.rgb = risk_color

        run = p.add_run(" / 100")
        run.font.name = Fonts.TITLE
        run.font.size = Pt(18)
        run.font.color.rgb = Colors.GRAY

        # Risk level label
        level_cell = gauge_table.rows[1].cells[0]
        self._set_cell_bg(level_cell, risk_color)
        p = level_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"  {risk_level}  ")
        run.font.name = Fonts.TITLE
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = Colors.WHITE

        self.doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════
        # QUICK STATS BOXES
        # ═══════════════════════════════════════════════════════════════════
        self._create_severity_boxes()

        self.doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════
        # REPORT METADATA TABLE
        # ═══════════════════════════════════════════════════════════════════
        meta_table = self.doc.add_table(rows=5, cols=4)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        meta_data = [
            ("Target", self.target_name, "Date", datetime.now().strftime('%B %d, %Y')),
            ("Report ID", f"SR-{datetime.now().strftime('%Y%m%d')}-001", "Version", "1.0"),
            ("Total Findings", str(self.stats.total_issues), "Filtered", str(len(self.filtered_vulns))),
            ("High Risk", str(self.stats.total_high), "Medium Risk", str(self.stats.total_medium)),
            ("Classification", "CONFIDENTIAL", "Distribution", "RESTRICTED"),
        ]

        for row_idx, (l1, v1, l2, v2) in enumerate(meta_data):
            row = meta_table.rows[row_idx]
            self._style_cell(row.cells[0], l1, bold=True, bg=Colors.LIGHT_GRAY)
            self._style_cell(row.cells[1], v1)
            self._style_cell(row.cells[2], l2, bold=True, bg=Colors.LIGHT_GRAY)
            self._style_cell(row.cells[3], v2)

        self._set_col_widths(meta_table, [1.2, 2, 1.2, 2])

        for _ in range(2):
            self.doc.add_paragraph()

        # ═══════════════════════════════════════════════════════════════════
        # FOOTER DISCLAIMER
        # ═══════════════════════════════════════════════════════════════════
        footer_table = self.doc.add_table(rows=1, cols=1)
        footer_cell = footer_table.rows[0].cells[0]
        self._set_cell_bg(footer_cell, Colors.LIGHT_GRAY)

        p = footer_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("⚠ CONFIDENTIAL - FOR AUTHORIZED PERSONNEL ONLY ⚠\n")
        run.font.name = Fonts.BODY
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = Colors.HIGH

        run = p.add_run("This document contains sensitive security information. Unauthorized disclosure is prohibited.")
        run.font.name = Fonts.BODY
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = Colors.DARK_GRAY

        self.doc.add_page_break()

    def _calculate_risk_score(self) -> Tuple[int, str, RGBColor]:
        """Calculate enterprise risk score (0-100) based on findings."""
        # Weighted scoring: Critical×25, High×15, Medium×5, Low×1, Info×0.1
        score = 0
        score += self.stats.total_high * 15  # High severity weight
        score += self.stats.total_medium * 5  # Medium severity weight
        score += self.stats.total_low * 1     # Low severity weight
        score += self.stats.total_info * 0.1  # Info weight

        # Add confidence multipliers
        score += self.stats.high_certain * 10  # Certain high findings add more
        score += self.stats.medium_certain * 3  # Certain medium findings

        # Normalize to 0-100 scale (cap at 100)
        risk_score = min(100, int(score))

        # Determine risk level
        if risk_score >= 75:
            return risk_score, "CRITICAL RISK", Colors.RISK_CRITICAL
        elif risk_score >= 50:
            return risk_score, "HIGH RISK", Colors.RISK_HIGH
        elif risk_score >= 25:
            return risk_score, "MEDIUM RISK", Colors.RISK_MEDIUM
        else:
            return risk_score, "LOW RISK", Colors.RISK_LOW

    def _create_severity_boxes(self):
        """Create premium visual severity summary boxes with icons."""
        table = self.doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Severity data with icons
        severities = [
            ("⬤", "HIGH", self.stats.total_high, Colors.HIGH, Colors.HIGH_BG),
            ("◆", "MEDIUM", self.stats.total_medium, Colors.MEDIUM, Colors.MEDIUM_BG),
            ("▲", "LOW", self.stats.total_low, Colors.LOW, Colors.LOW_BG),
            ("●", "INFO", self.stats.total_info, Colors.INFO, Colors.INFO_BG),
        ]

        for i, (icon, label, count, text_color, bg_color) in enumerate(severities):
            cell = table.rows[0].cells[i]
            cell.width = Inches(1.6)
            self._set_cell_bg(cell, bg_color)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # Icon
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(icon)
            run.font.size = Pt(14)
            run.font.color.rgb = text_color

            # Count
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(count))
            run.font.name = Fonts.TITLE
            run.font.size = Pt(36)
            run.font.bold = True
            run.font.color.rgb = text_color

            # Label
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.font.name = Fonts.BODY
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = text_color

            # Percentage of total
            total = self.stats.total_issues
            if total > 0:
                pct = (count / total) * 100
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"({pct:.1f}%)")
                run.font.name = Fonts.BODY
                run.font.size = Pt(8)
                run.font.color.rgb = Colors.GRAY

    def _create_executive_dashboard(self):
        """Create executive dashboard page."""
        self._add_section_title("1. EXECUTIVE SUMMARY")

        # Collect unique hosts/URLs for scope overview
        unique_hosts = set()
        for vuln in self.vulns:
            for inst in vuln.instances:
                # Extract host from various sources
                host = None
                if inst.host:
                    host = inst.host
                if inst.url:
                    try:
                        parsed = urlparse(inst.url)
                        if parsed.netloc:
                            host = parsed.netloc
                        elif inst.url.startswith(('http://', 'https://')):
                            # Handle malformed URLs
                            url_without_proto = inst.url.replace('http://', '').replace('https://', '')
                            if '/' in url_without_proto:
                                host = url_without_proto.split('/')[0]
                            elif url_without_proto:
                                host = url_without_proto
                    except (ValueError, AttributeError):
                        pass

                # Normalize: remove protocol prefixes if present
                if host:
                    host = host.replace('http://', '').replace('https://', '').strip()
                    if host:
                        unique_hosts.add(host)

        # Overview paragraph
        total_instances = sum(len(v.instances) for v in self.vulns)

        overview = f"""This report presents the findings of a comprehensive security assessment conducted on {self.target_name} using Burp Suite Professional scanner.

The assessment identified a total of {self.stats.total_issues} security findings across {len(self.vulns)} unique vulnerability types, affecting {total_instances} endpoints."""

        p = self.doc.add_paragraph(overview)
        p.paragraph_format.space_after = Pt(12)

        # Scope - Unique Hosts/URLs Tested
        if unique_hosts:
            self._add_subsection_title("1.1 Assessment Scope - Tested Hosts")

            p = self.doc.add_paragraph()
            run = p.add_run(f"The following {len(unique_hosts)} unique host(s) were identified during the assessment:")
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)

            # Create table for hosts
            hosts_per_row = 2
            sorted_hosts = sorted(unique_hosts)
            num_rows = (len(sorted_hosts) + hosts_per_row - 1) // hosts_per_row

            host_table = self.doc.add_table(rows=num_rows, cols=hosts_per_row)
            host_table.style = 'Table Grid'

            for idx, host in enumerate(sorted_hosts):
                row_idx = idx // hosts_per_row
                col_idx = idx % hosts_per_row
                cell = host_table.rows[row_idx].cells[col_idx]
                self._style_cell(cell, host, bg=Colors.LIGHT_GRAY)

            # Handle odd number of hosts - style empty cell
            if len(sorted_hosts) % hosts_per_row != 0:
                last_row = host_table.rows[num_rows - 1]
                empty_cell = last_row.cells[hosts_per_row - 1]
                self._style_cell(empty_cell, "", bg=Colors.LIGHT_GRAY)

            self._set_col_widths(host_table, [3.25, 3.25])
            self.doc.add_paragraph()

        # Key findings table
        self._add_subsection_title("1.2 Key Findings Overview")

        table = self.doc.add_table(rows=5, cols=5)
        table.style = 'Table Grid'

        # Header
        headers = ["Severity", "Certain", "Firm", "Tentative", "Total"]
        for i, h in enumerate(headers):
            self._style_cell(table.rows[0].cells[i], h, bold=True, bg=Colors.PRIMARY, text_color=Colors.WHITE)

        # Data rows
        data = [
            ("High", self.stats.high_certain, self.stats.high_firm, self.stats.high_tentative, self.stats.total_high, Colors.HIGH_BG, Colors.HIGH),
            ("Medium", self.stats.medium_certain, self.stats.medium_firm, self.stats.medium_tentative, self.stats.total_medium, Colors.MEDIUM_BG, Colors.MEDIUM),
            ("Low", self.stats.low_certain, self.stats.low_firm, self.stats.low_tentative, self.stats.total_low, Colors.LOW_BG, Colors.LOW),
            ("Information", self.stats.info_certain, self.stats.info_firm, self.stats.info_tentative, self.stats.total_info, Colors.INFO_BG, Colors.INFO),
        ]

        for row_idx, (sev, c, f, t, total, bg, text_col) in enumerate(data, 1):
            self._style_cell(table.rows[row_idx].cells[0], sev, bold=True, bg=bg, text_color=text_col)
            self._style_cell(table.rows[row_idx].cells[1], str(c), bg=bg)
            self._style_cell(table.rows[row_idx].cells[2], str(f), bg=bg)
            self._style_cell(table.rows[row_idx].cells[3], str(t), bg=bg)
            self._style_cell(table.rows[row_idx].cells[4], str(total), bold=True, bg=bg, text_color=text_col)

        self._set_col_widths(table, [1.5, 1, 1, 1, 1])

        self.doc.add_paragraph()

        # Recommendations summary
        self._add_subsection_title("1.3 Immediate Action Required")

        if self.stats.total_high > 0:
            p = self.doc.add_paragraph()
            run = p.add_run("CRITICAL: ")
            run.font.bold = True
            run.font.color.rgb = Colors.HIGH
            run = p.add_run(f"{self.stats.total_high} high-severity vulnerabilities require immediate remediation.")

        if self.stats.total_medium > 0:
            p = self.doc.add_paragraph()
            run = p.add_run("IMPORTANT: ")
            run.font.bold = True
            run.font.color.rgb = Colors.MEDIUM
            run = p.add_run(f"{self.stats.total_medium} medium-severity issues should be addressed promptly.")

        self.doc.add_page_break()

    def _create_visual_charts(self):
        """Create visual bar charts for severity distribution."""
        self._add_section_title("2. SEVERITY DISTRIBUTION CHARTS")

        p = self.doc.add_paragraph()
        p.add_run("The following charts illustrate the distribution of security findings:")
        p.paragraph_format.space_after = Pt(12)

        # ---- SEVERITY BAR CHART ----
        self._add_subsection_title("2.1 Findings by Severity Level")

        severity_data = [
            ("High", self.stats.total_high, Colors.HIGH, Colors.HIGH_BG),
            ("Medium", self.stats.total_medium, Colors.MEDIUM, Colors.MEDIUM_BG),
            ("Low", self.stats.total_low, Colors.LOW, Colors.LOW_BG),
            ("Information", self.stats.total_info, Colors.INFO, Colors.INFO_BG),
        ]

        max_count = max(d[1] for d in severity_data) if any(d[1] > 0 for d in severity_data) else 1
        bar_max_width = 20  # Number of cells for max bar width

        # Create bar chart table
        chart_table = self.doc.add_table(rows=4, cols=3)
        chart_table.style = 'Table Grid'

        for row_idx, (label, count, text_color, _bg_color) in enumerate(severity_data):
            row = chart_table.rows[row_idx]

            # Label cell
            self._style_cell(row.cells[0], label, bold=True, text_color=text_color)
            row.cells[0].width = Inches(1.2)

            # Bar cell - create visual bar using colored text blocks
            bar_cell = row.cells[1]
            bar_cell.width = Inches(4)

            bar_width = int((count / max_count) * bar_max_width) if max_count > 0 else 0
            bar_width = max(1, bar_width) if count > 0 else 0

            if bar_width > 0:
                p = bar_cell.paragraphs[0]
                run = p.add_run("█" * bar_width)
                run.font.size = Pt(12)
                run.font.color.rgb = text_color

            # Count cell
            self._style_cell(row.cells[2], str(count), bold=True, text_color=text_color, center=True)
            row.cells[2].width = Inches(0.8)

        self.doc.add_paragraph()

        # ---- CONFIDENCE DISTRIBUTION ----
        self._add_subsection_title("2.2 Findings by Confidence Level")

        confidence_data = [
            ("Certain", self.stats.high_certain + self.stats.medium_certain +
             self.stats.low_certain + self.stats.info_certain, Colors.SUCCESS),
            ("Firm", self.stats.high_firm + self.stats.medium_firm +
             self.stats.low_firm + self.stats.info_firm, Colors.MEDIUM),
            ("Tentative", self.stats.high_tentative + self.stats.medium_tentative +
             self.stats.low_tentative + self.stats.info_tentative, Colors.INFO),
        ]

        max_conf = max(d[1] for d in confidence_data) if any(d[1] > 0 for d in confidence_data) else 1

        conf_table = self.doc.add_table(rows=3, cols=3)
        conf_table.style = 'Table Grid'

        for row_idx, (label, count, bar_color) in enumerate(confidence_data):
            row = conf_table.rows[row_idx]

            self._style_cell(row.cells[0], label, bold=True)
            row.cells[0].width = Inches(1.2)

            bar_cell = row.cells[1]
            bar_cell.width = Inches(4)
            bar_width = int((count / max_conf) * bar_max_width) if max_conf > 0 else 0
            bar_width = max(1, bar_width) if count > 0 else 0

            if bar_width > 0:
                p = bar_cell.paragraphs[0]
                run = p.add_run("█" * bar_width)
                run.font.size = Pt(12)
                run.font.color.rgb = bar_color

            self._style_cell(row.cells[2], str(count), bold=True, center=True)
            row.cells[2].width = Inches(0.8)

        self.doc.add_paragraph()

        # ---- FILTERED SUMMARY (if filter applied) ----
        if self.severity_filter != {'High', 'Medium', 'Low', 'Information'}:
            self._add_subsection_title("2.3 Filtered Report Summary")

            p = self.doc.add_paragraph()
            run = p.add_run("Active Severity Filter: ")
            run.font.bold = True
            filter_text = ", ".join(sorted(self.severity_filter, key=lambda x: self.SEVERITY_ORDER.get(x, 4)))
            run = p.add_run(filter_text)
            run.font.color.rgb = Colors.PRIMARY

            # Filtered counts
            filtered_instances = sum(len(v.instances) for v in self.filtered_vulns)
            total_instances = sum(len(v.instances) for v in self.vulns)

            filter_table = self.doc.add_table(rows=3, cols=2)
            filter_table.style = 'Table Grid'

            self._style_cell(filter_table.rows[0].cells[0], "Vulnerability Types (Filtered)", bold=True, bg=Colors.LIGHT_GRAY)
            self._style_cell(filter_table.rows[0].cells[1], f"{len(self.filtered_vulns)} of {len(self.vulns)}")

            self._style_cell(filter_table.rows[1].cells[0], "Instances (Filtered)", bold=True, bg=Colors.LIGHT_GRAY)
            self._style_cell(filter_table.rows[1].cells[1], f"{filtered_instances} of {total_instances}")

            self._style_cell(filter_table.rows[2].cells[0], "Excluded Severities", bold=True, bg=Colors.LIGHT_GRAY)
            excluded = {'High', 'Medium', 'Low', 'Information'} - self.severity_filter
            excluded_text = ", ".join(sorted(excluded, key=lambda x: self.SEVERITY_ORDER.get(x, 4))) or "None"
            self._style_cell(filter_table.rows[2].cells[1], excluded_text)

            self._set_col_widths(filter_table, [2.5, 3])

        self.doc.add_page_break()

    def _create_owasp_compliance(self):
        """Create OWASP Top 10 2021 compliance mapping section."""
        self._add_section_title("3. OWASP TOP 10 COMPLIANCE MAPPING")

        p = self.doc.add_paragraph()
        run = p.add_run("This section maps identified vulnerabilities to the OWASP Top 10 2021 categories:")
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)

        # Count vulnerabilities by OWASP category
        owasp_counts = Counter()
        for vuln in self.filtered_vulns:
            owasp_code, owasp_name = get_owasp_category(vuln.title)
            owasp_counts[f"{owasp_code} - {owasp_name}"] += len(vuln.instances)

        # Create OWASP mapping table
        owasp_table = self.doc.add_table(rows=1, cols=4)
        owasp_table.style = 'Table Grid'

        headers = ["OWASP Category", "Findings", "Risk Level", "Status"]
        for i, h in enumerate(headers):
            self._style_cell(owasp_table.rows[0].cells[i], h, bold=True,
                           bg=Colors.PRIMARY, text_color=Colors.WHITE, center=True)

        # All OWASP Top 10 categories
        all_owasp = [
            ("A01:2021", "Broken Access Control"),
            ("A02:2021", "Cryptographic Failures"),
            ("A03:2021", "Injection"),
            ("A04:2021", "Insecure Design"),
            ("A05:2021", "Security Misconfiguration"),
            ("A06:2021", "Vulnerable Components"),
            ("A07:2021", "Auth Failures"),
            ("A08:2021", "Software/Data Integrity"),
            ("A09:2021", "Logging Failures"),
            ("A10:2021", "SSRF"),
        ]

        for code, name in all_owasp:
            key = f"{code} - {name}"
            count = owasp_counts.get(key, 0)
            row = owasp_table.add_row()

            self._style_cell(row.cells[0], f"{code}\n{name}", bold=True)

            if count > 0:
                self._style_cell(row.cells[1], str(count), center=True, bold=True,
                               text_color=Colors.HIGH if count > 10 else Colors.MEDIUM if count > 3 else Colors.LOW)

                # Risk level based on count
                if count > 10:
                    risk = "HIGH"
                    risk_color = Colors.HIGH
                elif count > 3:
                    risk = "MEDIUM"
                    risk_color = Colors.MEDIUM
                else:
                    risk = "LOW"
                    risk_color = Colors.LOW
                self._style_cell(row.cells[2], risk, center=True, bold=True, text_color=risk_color)
                self._style_cell(row.cells[3], "⚠ ACTION NEEDED", center=True, text_color=Colors.HIGH)
            else:
                self._style_cell(row.cells[1], "0", center=True, text_color=Colors.GRAY)
                self._style_cell(row.cells[2], "N/A", center=True, text_color=Colors.GRAY)
                self._style_cell(row.cells[3], "✓ COMPLIANT", center=True, text_color=Colors.SUCCESS)

        self._set_col_widths(owasp_table, [2.5, 1, 1, 1.5])
        self._style_table_borders(owasp_table)

        self.doc.add_paragraph()

        # Summary
        affected = sum(1 for c, n in all_owasp if owasp_counts.get(f"{c} - {n}", 0) > 0)
        p = self.doc.add_paragraph()
        run = p.add_run(f"Summary: ")
        run.font.bold = True
        run = p.add_run(f"{affected} of 10 OWASP categories have identified findings. ")
        if affected >= 5:
            run = p.add_run("Immediate attention required.")
            run.font.color.rgb = Colors.HIGH
            run.font.bold = True
        elif affected >= 3:
            run = p.add_run("Review and prioritize remediation.")
            run.font.color.rgb = Colors.MEDIUM
        else:
            run = p.add_run("Good security posture, continue monitoring.")
            run.font.color.rgb = Colors.SUCCESS

        self.doc.add_page_break()

    def _create_remediation_priority(self):
        """Create remediation priority matrix with effort estimation."""
        self._add_section_title("4. REMEDIATION PRIORITY MATRIX")

        p = self.doc.add_paragraph()
        run = p.add_run("The following matrix prioritizes vulnerabilities based on risk and remediation effort:")
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)

        # Priority table
        priority_table = self.doc.add_table(rows=1, cols=5)
        priority_table.style = 'Table Grid'

        headers = ["Priority", "Vulnerability Type", "Severity", "Instances", "Recommended Action"]
        for i, h in enumerate(headers):
            self._style_cell(priority_table.rows[0].cells[i], h, bold=True,
                           bg=Colors.PRIMARY, text_color=Colors.WHITE, center=True)

        # Sort by severity and instance count
        sorted_vulns = sorted(self.filtered_vulns,
                             key=lambda v: (self.SEVERITY_ORDER.get(v.severity, 4), -len(v.instances)))

        priority = 0
        for vuln in sorted_vulns[:15]:  # Top 15 priorities
            priority += 1
            row = priority_table.add_row()

            # Priority number with color coding
            if priority <= 3:
                prio_color = Colors.HIGH
                action = "Immediate remediation required"
            elif priority <= 7:
                prio_color = Colors.MEDIUM
                action = "Address within 30 days"
            else:
                prio_color = Colors.LOW
                action = "Schedule for next cycle"

            self._style_cell(row.cells[0], f"P{priority}", bold=True,
                           text_color=prio_color, center=True)

            title = vuln.title[:40] + '...' if len(vuln.title) > 40 else vuln.title
            self._style_cell(row.cells[1], title)

            sev_color = self._get_severity_color(vuln.severity)
            self._style_cell(row.cells[2], vuln.severity, center=True,
                           text_color=sev_color, bold=True)

            self._style_cell(row.cells[3], str(len(vuln.instances)), center=True)
            self._style_cell(row.cells[4], action, text_color=prio_color)

        self._set_col_widths(priority_table, [0.6, 2.2, 0.8, 0.7, 2])
        self._style_table_borders(priority_table)

        if len(self.filtered_vulns) > 15:
            p = self.doc.add_paragraph()
            run = p.add_run(f"Note: {len(self.filtered_vulns) - 15} additional findings not shown. See detailed findings section.")
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = Colors.GRAY

        self.doc.add_page_break()

    def _create_risk_matrix(self):
        """Create risk rating matrix."""
        self._add_section_title("5. RISK ASSESSMENT MATRIX")

        p = self.doc.add_paragraph()
        p.add_run("The following matrix illustrates the distribution of findings by severity and confidence levels:")
        p.paragraph_format.space_after = Pt(12)

        # Matrix table
        table = self.doc.add_table(rows=5, cols=4)
        table.style = 'Table Grid'

        # Headers
        matrix_headers = ["", "Certain", "Firm", "Tentative"]
        for i, h in enumerate(matrix_headers):
            self._style_cell(table.rows[0].cells[i], h, bold=True, bg=Colors.PRIMARY, text_color=Colors.WHITE, center=True)

        # Rows
        matrix_data = [
            ("HIGH", self.stats.high_certain, self.stats.high_firm, self.stats.high_tentative, Colors.HIGH, Colors.HIGH_BG),
            ("MEDIUM", self.stats.medium_certain, self.stats.medium_firm, self.stats.medium_tentative, Colors.MEDIUM, Colors.MEDIUM_BG),
            ("LOW", self.stats.low_certain, self.stats.low_firm, self.stats.low_tentative, Colors.LOW, Colors.LOW_BG),
            ("INFO", self.stats.info_certain, self.stats.info_firm, self.stats.info_tentative, Colors.INFO, Colors.INFO_BG),
        ]

        for row_idx, (label, c, f, t, text_col, bg) in enumerate(matrix_data, 1):
            self._style_cell(table.rows[row_idx].cells[0], label, bold=True, bg=bg, text_color=text_col, center=True)

            # Color cells based on count
            for col_idx, count in enumerate([c, f, t], 1):
                if count > 0:
                    self._style_cell(table.rows[row_idx].cells[col_idx], str(count), bg=bg, text_color=text_col, center=True, bold=True)
                else:
                    self._style_cell(table.rows[row_idx].cells[col_idx], "-", center=True, text_color=Colors.GRAY)

        self._set_col_widths(table, [1.5, 1.5, 1.5, 1.5])

        self.doc.add_paragraph()
        self.doc.add_page_break()

    def _create_findings_summary(self):
        """Create findings summary grouped by domain."""
        self._add_section_title("6. VULNERABILITY SUMMARY")

        # Show filter info
        filter_text = ", ".join(sorted(self.severity_filter, key=lambda x: self.SEVERITY_ORDER.get(x, 4)))
        p = self.doc.add_paragraph()
        run = p.add_run(f"Showing: {filter_text} severity findings ({len(self.filtered_vulns)} vulnerability types)")
        run.font.italic = True
        run.font.color.rgb = Colors.GRAY
        p.paragraph_format.space_after = Pt(12)

        # Handle empty filter results
        if not self.filtered_vulns:
            p = self.doc.add_paragraph()
            run = p.add_run("No vulnerabilities found matching the selected severity filter.")
            run.font.italic = True
            run.font.color.rgb = Colors.GRAY
            self.doc.add_page_break()
            return

        # Summary table - uses pre-filtered and sorted self.filtered_vulns
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        headers = ["#", "Vulnerability", "Severity", "Confidence", "Instances"]
        for i, h in enumerate(headers):
            self._style_cell(table.rows[0].cells[i], h, bold=True, bg=Colors.PRIMARY, text_color=Colors.WHITE)

        for idx, vuln in enumerate(self.filtered_vulns, 1):
            row = table.add_row()
            self._style_cell(row.cells[0], str(idx), center=True)
            self._style_cell(row.cells[1], vuln.title[:55] + ('...' if len(vuln.title) > 55 else ''))

            sev_color = self._get_severity_color(vuln.severity)
            self._style_cell(row.cells[2], vuln.severity, bold=True, text_color=sev_color, center=True)
            self._style_cell(row.cells[3], vuln.confidence, center=True)
            self._style_cell(row.cells[4], str(len(vuln.instances)), center=True)

        self._set_col_widths(table, [0.4, 3.5, 0.8, 0.8, 0.7])

        self.doc.add_page_break()

    def _create_detailed_findings(self):
        """Create detailed findings section - uses filtered vulnerabilities."""
        self._add_section_title("7. DETAILED FINDINGS")

        # Handle case where no vulnerabilities match the filter
        if not self.filtered_vulns:
            p = self.doc.add_paragraph()
            run = p.add_run("No vulnerabilities match the selected severity filter.")
            run.font.italic = True
            run.font.color.rgb = Colors.GRAY
            p.paragraph_format.space_before = Pt(12)

            p = self.doc.add_paragraph()
            filter_text = ", ".join(sorted(self.severity_filter, key=lambda x: self.SEVERITY_ORDER.get(x, 4)))
            run = p.add_run(f"Current filter: {filter_text}")
            run.font.size = Pt(10)
            run.font.color.rgb = Colors.GRAY
            return

        # Uses pre-filtered self.filtered_vulns
        for idx, vuln in enumerate(self.filtered_vulns, 1):
            if idx % 5 == 0:
                logger.info(f"Writing finding {idx}/{len(self.filtered_vulns)}...")
            self._create_vulnerability_section(idx, vuln)

    def _create_vulnerability_section(self, idx: int, vuln: Vulnerability):
        """Create a single vulnerability section."""
        # Title with severity badge
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)

        sev_color = self._get_severity_color(vuln.severity)
        run = p.add_run(f"[{vuln.severity.upper()}] ")
        run.font.bold = True
        run.font.color.rgb = sev_color
        run.font.size = Pt(12)

        run = p.add_run(f"7.{idx} {vuln.title}")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = Colors.PRIMARY

        # Metadata table
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        self._style_cell(table.rows[0].cells[0], "Severity", bold=True, bg=Colors.LIGHT_GRAY)
        self._style_cell(table.rows[0].cells[1], vuln.severity, bold=True, text_color=sev_color)
        self._style_cell(table.rows[0].cells[2], "Confidence", bold=True, bg=Colors.LIGHT_GRAY)
        self._style_cell(table.rows[0].cells[3], vuln.confidence)

        self._style_cell(table.rows[1].cells[0], "Instances", bold=True, bg=Colors.LIGHT_GRAY)
        self._style_cell(table.rows[1].cells[1], str(len(vuln.instances)))
        self._style_cell(table.rows[1].cells[2], "CWE", bold=True, bg=Colors.LIGHT_GRAY)
        cwe = vuln.cwe_classifications[0][0] if vuln.cwe_classifications else "N/A"
        cwe = cwe[:37] + '...' if len(cwe) > 40 else cwe
        self._style_cell(table.rows[1].cells[3], cwe)

        self._set_col_widths(table, [1, 2, 1, 2.5])
        self.doc.add_paragraph()

        # Issue Background
        if vuln.issue_background:
            self._add_mini_header("Issue Background")
            p = self.doc.add_paragraph(vuln.issue_background)
            p.paragraph_format.space_after = Pt(6)

        # Remediation
        if vuln.issue_remediation:
            self._add_mini_header("Remediation")
            p = self.doc.add_paragraph(vuln.issue_remediation)
            p.paragraph_format.space_after = Pt(6)

        # References
        if vuln.references:
            self._add_mini_header("References")
            for ref_text, _ref_url in vuln.references[:5]:
                p = self.doc.add_paragraph()
                run = p.add_run(f"  - {ref_text}")
                run.font.size = Pt(9)

        # CWE Classifications
        if vuln.cwe_classifications:
            self._add_mini_header("Vulnerability Classifications")
            for cwe_name, _ in vuln.cwe_classifications:
                p = self.doc.add_paragraph()
                run = p.add_run(f"  - {cwe_name}")
                run.font.size = Pt(9)

        # Group instances by domain
        self._add_mini_header(f"Affected Endpoints ({len(vuln.instances)} instances)")

        # Group by domain
        domains: Dict[str, List[VulnerabilityInstance]] = defaultdict(list)
        for inst in vuln.instances:
            domain = "Unknown"
            try:
                if inst.url:
                    parsed = urlparse(inst.url)
                    domain = parsed.netloc or inst.host or "Unknown"
                elif inst.host:
                    domain = inst.host
            except (ValueError, AttributeError):
                domain = inst.host or "Unknown"
            domains[domain].append(inst)

        # Display grouped by domain
        for domain, instances in domains.items():
            p = self.doc.add_paragraph()
            run = p.add_run(f"Domain: {domain} ({len(instances)} endpoints)")
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = Colors.PRIMARY

            for inst_idx, instance in enumerate(instances, 1):
                self._create_instance_details(inst_idx, instance, compact=False)

        self.doc.add_page_break()

    def _create_instance_details(self, idx: int, instance: VulnerabilityInstance, compact: bool = False):
        """Create instance details."""
        # Instance header
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"{idx}. ")
        run.font.bold = True
        run.font.size = Pt(9)

        path_display = instance.path or instance.url or "[No URL available]"
        if len(path_display) > 80:
            path_display = path_display[:77] + "..."
        run = p.add_run(path_display)
        run.font.size = Pt(9)
        run.font.name = Fonts.CODE

        # Issue detail
        if instance.issue_detail and not compact:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            detail_text = instance.issue_detail
            if len(detail_text) > 500:
                detail_text = detail_text[:497] + '...'
            run = p.add_run(detail_text)
            run.font.size = Pt(8)
            run.font.italic = True

        # Evidence-only mode: show only highlighted portions or smart summary
        if self.evidence_only:
            has_evidence = False

            # Show request evidence only
            if instance.request_highlights:
                self._add_evidence_block("Request Evidence", instance.request_highlights)
                has_evidence = True
            elif instance.request:
                # Show smart summary of request (first line with method/URL)
                self._add_request_summary(instance.request)
                has_evidence = True

            # Show response evidence only
            if instance.response_highlights:
                self._add_evidence_block("Response Evidence", instance.response_highlights)
                has_evidence = True
            elif instance.response:
                # Show smart summary of response (status line + key headers)
                self._add_response_summary(instance.response)
                has_evidence = True

            # If nothing at all, show a note
            if not has_evidence:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run("[Vulnerability detected by scanner analysis]")
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = Colors.GRAY
        else:
            # Full mode: show complete request/response
            if instance.request and not compact:
                self._add_code_block("Request", instance.request, instance.request_highlights)

            if instance.response and not compact:
                self._add_code_block("Response", instance.response, instance.response_highlights)

    def _add_request_summary(self, request: str):
        """Add smart summary of HTTP request (method, URL, key headers only)."""
        if not request:
            return

        lines = request.split('\n')
        summary_lines = []
        important_headers = ('host:', 'content-type:', 'cookie:', 'authorization:',
                           'origin:', 'referer:', 'x-forwarded', 'user-agent:')
        got_first_line = False

        for line in lines:
            line = line.strip()
            if not line:
                continue  # Skip empty lines
            if not got_first_line:  # First non-empty line: GET /path HTTP/1.1
                summary_lines.append(line)
                got_first_line = True
            elif line.lower().startswith(important_headers):
                summary_lines.append(line[:100] + '...' if len(line) > 100 else line)
            if len(summary_lines) >= 6:  # Limit to 6 key lines
                break

        if summary_lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run("Request Summary:")
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = Colors.PRIMARY

            table = self.doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.rows[0].cells[0]
            self._set_cell_bg(cell, Colors.CODE_BG)

            p = cell.paragraphs[0]
            for line in summary_lines:
                run = p.add_run(line + '\n')
                run.font.name = Fonts.CODE
                run.font.size = Pt(8)

    def _add_response_summary(self, response: str):
        """Add smart summary of HTTP response (status, key headers only)."""
        if not response:
            return

        lines = response.split('\n')
        summary_lines = []
        important_headers = ('content-type:', 'server:', 'set-cookie:', 'location:',
                           'x-', 'access-control', 'strict-transport', 'content-security')
        got_first_line = False

        for line in lines:
            line = line.strip()
            if not line:
                continue  # Skip empty lines
            if not got_first_line:  # First non-empty line: HTTP/1.1 200 OK
                summary_lines.append(line)
                got_first_line = True
            elif line.lower().startswith(important_headers):
                summary_lines.append(line[:100] + '...' if len(line) > 100 else line)
            if len(summary_lines) >= 6:  # Limit to 6 key lines
                break

        if summary_lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run("Response Summary:")
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = Colors.PRIMARY

            table = self.doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.rows[0].cells[0]
            self._set_cell_bg(cell, Colors.CODE_BG)

            p = cell.paragraphs[0]
            for line in summary_lines:
                run = p.add_run(line + '\n')
                run.font.name = Fonts.CODE
                run.font.size = Pt(8)

    def _add_evidence_block(self, title: str, highlights: List[str]):
        """Add evidence-only block with highlighted portions."""
        # Skip if no highlights
        if not highlights:
            return

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"{title}:")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = Colors.HIGH

        # Evidence table with red left border to indicate importance
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        self._set_cell_bg(cell, Colors.HIGH_BG)

        first = True
        for idx, evidence in enumerate(highlights):
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()

            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Evidence number
            if len(highlights) > 1:
                run = p.add_run(f"[{idx + 1}] ")
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = Colors.HIGH

            # Evidence content (truncate if too long with indicator)
            if len(evidence) > 500:
                evidence_text = evidence[:497] + '...'
            else:
                evidence_text = evidence
            run = p.add_run(evidence_text)
            run.font.name = Fonts.CODE
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = Colors.DARK_GRAY

        self.doc.add_paragraph()

    def _add_code_block(self, title: str, content: str, highlights: List[str] = None):
        """Add formatted code block."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(f"{title}:")
        run.font.bold = True
        run.font.size = Pt(9)

        # Code table
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        self._set_cell_bg(cell, Colors.CODE_BG)

        lines = content.split('\n')
        first = True

        for line in lines:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # Only highlight if the match is significant (>3 chars) to avoid false positives
            is_highlight = highlights and any(h in line for h in highlights if len(h) > 3)

            # Handle long lines with truncation indicator
            if len(line) > 400:
                display_line = line[:397] + '...'
            else:
                display_line = line
            run = p.add_run(display_line)
            run.font.name = Fonts.CODE
            run.font.size = Pt(7)

            if is_highlight:
                run.font.bold = True
                run.font.color.rgb = Colors.HIGH

    def _create_appendix(self):
        """Create appendix."""
        self._add_section_title("8. APPENDIX")

        # Severity definitions
        self._add_subsection_title("8.1 Severity Definitions")

        definitions = [
            ("HIGH", Colors.HIGH, "Critical vulnerabilities that could lead to system compromise, data breach, or significant business impact. Immediate remediation required."),
            ("MEDIUM", Colors.MEDIUM, "Significant vulnerabilities that could be exploited under certain conditions. Should be addressed within a short timeframe."),
            ("LOW", Colors.LOW, "Minor vulnerabilities with limited direct impact. Address as part of regular maintenance cycles."),
            ("INFORMATION", Colors.INFO, "Informational findings that may indicate areas for security hardening but do not represent exploitable vulnerabilities."),
        ]

        for sev, color, desc in definitions:
            p = self.doc.add_paragraph()
            run = p.add_run(f"{sev}: ")
            run.font.bold = True
            run.font.color.rgb = color
            run = p.add_run(desc)
            run.font.size = Pt(10)

        self.doc.add_paragraph()

        # Confidence definitions
        self._add_subsection_title("8.2 Confidence Levels")

        conf_defs = [
            ("Certain", "Vulnerability confirmed through reliable, deterministic detection techniques."),
            ("Firm", "High likelihood based on strong indicators and behavioral patterns."),
            ("Tentative", "Possible vulnerability that requires manual verification to confirm."),
        ]

        for level, desc in conf_defs:
            p = self.doc.add_paragraph()
            run = p.add_run(f"{level}: ")
            run.font.bold = True
            run = p.add_run(desc)
            run.font.size = Pt(10)

        self.doc.add_paragraph()

        # Methodology
        self._add_subsection_title("8.3 Assessment Methodology")

        methodology = """This assessment was conducted using Burp Suite Professional, an industry-leading web application security testing platform.

Testing Scope:
- Automated vulnerability scanning
- Analysis of application responses
- Detection of OWASP Top 10 vulnerabilities
- Security header and configuration analysis
- Authentication and session management testing

Limitations:
- Automated scanning may not identify all vulnerability types
- Business logic vulnerabilities require manual testing
- Results should be validated before remediation

Recommendations:
- Supplement automated scanning with manual penetration testing
- Implement a regular security assessment schedule
- Address findings based on risk prioritization"""

        for para in methodology.split('\n\n'):
            p = self.doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(8)

    # ========================================================================
    # NETWORK PENTEST SECTION
    # ========================================================================

    def _parse_network_reports(self):
        """Parse all Excel files from network reports folder."""
        if not self.network_reports_folder:
            return

        logger.info(f"Parsing network reports from: {self.network_reports_folder}")

        for filename in os.listdir(self.network_reports_folder):
            if not filename.endswith('.xlsx'):
                continue

            filepath = os.path.join(self.network_reports_folder, filename)
            try:
                # Read Excel with proper structure (skip first 2 rows which are headers)
                df = pd.read_excel(filepath, skiprows=2, header=0)

                # Validate column count before assigning names
                expected_cols = ['Service', 'Port', 'Version', 'CVE', 'CVSS', 'Severity', 'Exploits']
                if len(df.columns) < 7:
                    logger.warning(f"  Skipping {filename}: Expected 7 columns, found {len(df.columns)}")
                    continue
                elif len(df.columns) > 7:
                    # More columns than expected - just use first 7
                    df = df.iloc[:, :7]

                df.columns = expected_cols

                # Clean data
                df = df.dropna(subset=['CVE'])  # Remove rows without CVE
                df['CVSS'] = pd.to_numeric(df['CVSS'], errors='coerce')
                # Normalize severity to uppercase for consistent filtering
                df['Severity'] = df['Severity'].astype(str).str.upper().str.strip()

                # Extract domain name from filename
                domain = filename.replace('.xlsx', '').replace('-', '.')

                self.network_vulns[domain] = {
                    'filename': filename,
                    'data': df,
                    'total': len(df),
                    'critical': len(df[df['Severity'] == 'CRITICAL']),
                    'high': len(df[df['Severity'] == 'HIGH']),
                    'medium': len(df[df['Severity'] == 'MEDIUM']),
                    'low': len(df[df['Severity'] == 'LOW']),
                }

                logger.info(f"  Parsed {filename}: {len(df)} vulnerabilities")

            except Exception as e:
                logger.warning(f"  Error parsing {filename}: {e}")

        if len(self.network_vulns) == 0:
            logger.warning("No valid Excel files found in network reports folder")
        else:
            logger.info(f"Total network domains: {len(self.network_vulns)}")

    def _create_network_assessment_section(self):
        """Create Section 7: Network Vulnerability Assessment."""
        if not self.network_vulns:
            return

        self.doc.add_page_break()
        self._add_section_title("9. NETWORK VULNERABILITY ASSESSMENT")

        # Overview
        total_vulns = sum(d['total'] for d in self.network_vulns.values())
        total_critical = sum(d['critical'] for d in self.network_vulns.values())
        total_high = sum(d['high'] for d in self.network_vulns.values())
        total_medium = sum(d['medium'] for d in self.network_vulns.values())
        total_low = sum(d['low'] for d in self.network_vulns.values())

        p = self.doc.add_paragraph()
        p.add_run(f"Network vulnerability scanning identified ")
        run = p.add_run(f"{total_vulns} CVEs")
        run.font.bold = True
        p.add_run(f" across {len(self.network_vulns)} hosts/domains.")
        p.paragraph_format.space_after = Pt(12)

        # Summary table
        self._add_subsection_title("9.1 Network Vulnerability Summary")

        summary_table = self.doc.add_table(rows=1, cols=6)
        summary_table.style = 'Table Grid'

        # Headers
        headers = ["Host/Domain", "Critical", "High", "Medium", "Low", "Total"]
        for i, h in enumerate(headers):
            self._style_cell(summary_table.rows[0].cells[i], h, bold=True,
                           bg=Colors.PRIMARY, text_color=Colors.WHITE, center=True)

        # Data rows
        for domain, data in sorted(self.network_vulns.items()):
            row = summary_table.add_row()
            # Truncate long domain names
            domain_display = domain[:35] + '...' if len(domain) > 35 else domain
            self._style_cell(row.cells[0], domain_display)

            # Critical with red background
            if data['critical'] > 0:
                self._style_cell(row.cells[1], str(data['critical']), bold=True,
                               text_color=Colors.WHITE, bg=Colors.CRITICAL, center=True)
            else:
                self._style_cell(row.cells[1], str(data['critical']), center=True)

            # High
            if data['high'] > 0:
                self._style_cell(row.cells[2], str(data['high']), bold=True,
                               text_color=Colors.HIGH, center=True)
            else:
                self._style_cell(row.cells[2], str(data['high']), center=True)

            # Medium
            if data['medium'] > 0:
                self._style_cell(row.cells[3], str(data['medium']),
                               text_color=Colors.MEDIUM, center=True)
            else:
                self._style_cell(row.cells[3], str(data['medium']), center=True)

            # Low
            self._style_cell(row.cells[4], str(data['low']), center=True)

            # Total
            self._style_cell(row.cells[5], str(data['total']), bold=True, center=True)

        # Totals row
        total_row = summary_table.add_row()
        self._style_cell(total_row.cells[0], "TOTAL", bold=True, bg=Colors.LIGHT_GRAY)
        self._style_cell(total_row.cells[1], str(total_critical), bold=True,
                        bg=Colors.LIGHT_GRAY, center=True)
        self._style_cell(total_row.cells[2], str(total_high), bold=True,
                        bg=Colors.LIGHT_GRAY, center=True)
        self._style_cell(total_row.cells[3], str(total_medium), bold=True,
                        bg=Colors.LIGHT_GRAY, center=True)
        self._style_cell(total_row.cells[4], str(total_low), bold=True,
                        bg=Colors.LIGHT_GRAY, center=True)
        self._style_cell(total_row.cells[5], str(total_vulns), bold=True,
                        bg=Colors.LIGHT_GRAY, center=True)

        self._set_col_widths(summary_table, [2.5, 0.7, 0.7, 0.7, 0.7, 0.7])
        self._style_table_borders(summary_table)  # Apply professional border styling
        self.doc.add_paragraph()

        # Detailed findings by domain
        self._add_subsection_title("9.2 Detailed Network Findings by Host")

        domain_idx = 0
        for domain, data in sorted(self.network_vulns.items()):
            domain_idx += 1

            # Domain header
            p = self.doc.add_paragraph()
            run = p.add_run(f"9.2.{domain_idx} {domain}")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = Colors.PRIMARY
            p.paragraph_format.space_before = Pt(12)

            # Stats line
            p = self.doc.add_paragraph()
            run = p.add_run(f"Critical: {data['critical']} | High: {data['high']} | Medium: {data['medium']} | Low: {data['low']}")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = Colors.GRAY

            # CVE table for this domain
            df = data['data']

            # Group by severity for organized display
            for severity in ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW']:
                sev_df = df[df['Severity'] == severity]
                if len(sev_df) == 0:
                    continue

                # Create table for this severity
                vuln_table = self.doc.add_table(rows=1, cols=5)
                vuln_table.style = 'Table Grid'

                # Header with severity color
                sev_color = {
                    'CRITICAL': Colors.CRITICAL,
                    'HIGH': Colors.HIGH,
                    'MEDIUM': Colors.MEDIUM,
                    'LOW': Colors.LOW
                }.get(severity, Colors.INFO)

                sev_bg = {
                    'CRITICAL': Colors.HIGH_BG,
                    'HIGH': Colors.HIGH_BG,
                    'MEDIUM': Colors.MEDIUM_BG,
                    'LOW': Colors.LOW_BG
                }.get(severity, Colors.INFO_BG)

                headers = [severity, "Port", "Version", "CVE", "CVSS"]
                for i, h in enumerate(headers):
                    self._style_cell(vuln_table.rows[0].cells[i], h, bold=True,
                                   bg=sev_bg, text_color=sev_color)

                # Add ALL vulnerability rows (no limit - show complete list)
                for _, vuln_row in sev_df.iterrows():
                    row = vuln_table.add_row()
                    # Handle NaN values properly - convert to string safely
                    service = str(vuln_row['Service'])[:15] if pd.notna(vuln_row['Service']) else ''
                    self._style_cell(row.cells[0], service)
                    # Port: convert to int to avoid "443.0" display, handle non-numeric
                    port_val = vuln_row['Port']
                    try:
                        port_str = str(int(float(port_val))) if pd.notna(port_val) else ''
                    except (ValueError, TypeError):
                        port_str = str(port_val) if pd.notna(port_val) else ''
                    self._style_cell(row.cells[1], port_str, center=True)
                    version = str(vuln_row['Version'])[:25] if pd.notna(vuln_row['Version']) else ''
                    self._style_cell(row.cells[2], version)
                    cve = str(vuln_row['CVE']) if pd.notna(vuln_row['CVE']) else ''
                    self._style_cell(row.cells[3], cve)
                    # CVSS: safely convert and check, handle non-numeric values
                    cvss_val = vuln_row['CVSS']
                    try:
                        cvss_num = float(cvss_val) if pd.notna(cvss_val) else 0
                        cvss = f"{cvss_num:.1f}"
                    except (ValueError, TypeError):
                        cvss_num = 0
                        cvss = str(cvss_val) if pd.notna(cvss_val) else ''
                    self._style_cell(row.cells[4], cvss, center=True,
                                   text_color=sev_color if cvss_num >= 7.0 else None)

                self._set_col_widths(vuln_table, [1.0, 0.6, 1.8, 1.8, 0.6])
                self._style_table_borders(vuln_table)  # Apply professional border styling

            self.doc.add_paragraph()  # Space between domains

            # Add page break every 3 domains
            if domain_idx % 3 == 0 and domain_idx < len(self.network_vulns):
                self.doc.add_page_break()

    # ========================================================================
    # HELPER METHODS
    # ========================================================================

    def _add_section_title(self, text: str):
        """Add main section title with professional styling."""
        # Title paragraph with left accent bar
        p = self.doc.add_paragraph()

        # Add thick left border as accent
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left_bdr = OxmlElement('w:left')
        left_bdr.set(qn('w:val'), 'single')
        left_bdr.set(qn('w:sz'), '36')  # Thick left border
        left_bdr.set(qn('w:space'), '8')
        left_bdr.set(qn('w:color'), '003366')  # Primary color
        pBdr.append(left_bdr)

        # Add bottom border
        bottom_bdr = OxmlElement('w:bottom')
        bottom_bdr.set(qn('w:val'), 'single')
        bottom_bdr.set(qn('w:sz'), '6')
        bottom_bdr.set(qn('w:space'), '1')
        bottom_bdr.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom_bdr)
        pPr.append(pBdr)

        run = p.add_run(text)
        run.font.name = Fonts.HEADING
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = Colors.PRIMARY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.left_indent = Inches(0.15)

    def _add_subsection_title(self, text: str):
        """Add subsection title with left border accent."""
        p = self.doc.add_paragraph()

        # Add left border accent
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left_bdr = OxmlElement('w:left')
        left_bdr.set(qn('w:val'), 'single')
        left_bdr.set(qn('w:sz'), '24')  # Thick border
        left_bdr.set(qn('w:space'), '4')
        left_bdr.set(qn('w:color'), '003366')  # Primary color
        pBdr.append(left_bdr)
        pPr.append(pBdr)

        run = p.add_run(text)
        run.font.name = Fonts.HEADING
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = Colors.SECONDARY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)

    def _add_mini_header(self, text: str):
        """Add mini header."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = Fonts.BODY
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = Colors.DARK_GRAY
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    def _style_cell(self, cell, text: str, bold: bool = False, bg: RGBColor = None,
                   text_color: RGBColor = None, center: bool = False):
        """Style table cell."""
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if p.runs:
            run = p.runs[0]
        else:
            p.clear()
            run = p.add_run(text)

        run.font.name = Fonts.BODY
        run.font.size = Pt(9)
        if bold:
            run.font.bold = True
        if text_color:
            run.font.color.rgb = text_color
        if bg:
            self._set_cell_bg(cell, bg)

    def _style_info_table(self, table, data: List[Tuple[str, str]]):
        """Style info table."""
        for i, (label, value) in enumerate(data):
            self._style_cell(table.rows[i].cells[0], label, bold=True, bg=Colors.LIGHT_GRAY)
            self._style_cell(table.rows[i].cells[1], value)
        self._set_col_widths(table, [2, 3])

    def _set_cell_bg(self, cell, color: RGBColor):
        """Set cell background."""
        hex_color = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_col_widths(self, table: Table, widths: List[float]):
        """Set column widths."""
        for i, width in enumerate(widths):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    def _style_table_borders(self, table: Table):
        """Apply professional border styling to table."""
        tbl = table._tbl
        # Get existing tblPr or create and attach a new one
        if tbl.tblPr is not None:
            tblPr = tbl.tblPr
        else:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)  # Attach to table element

        # Table borders
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CCCCCC')
            tblBorders.append(border)

        tblPr.append(tblBorders)

        # Add shadow effect (cell margins for depth appearance)
        tblCellMar = OxmlElement('w:tblCellMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '80')
            margin.set(qn('w:type'), 'dxa')
            tblCellMar.append(margin)
        tblPr.append(tblCellMar)

    def _get_severity_color(self, severity: str) -> RGBColor:
        """Get severity color."""
        return {
            'High': Colors.HIGH,
            'Medium': Colors.MEDIUM,
            'Low': Colors.LOW,
            'Information': Colors.INFO,
        }.get(severity, Colors.INFO)

    def _get_risk_level(self) -> Tuple[str, RGBColor]:
        """Determine overall risk level."""
        if self.stats.total_high > 5:
            return "CRITICAL", Colors.CRITICAL
        elif self.stats.total_high > 0:
            return "HIGH", Colors.HIGH
        elif self.stats.total_medium > 5:
            return "MEDIUM-HIGH", Colors.MEDIUM
        elif self.stats.total_medium > 0:
            return "MEDIUM", Colors.MEDIUM
        elif self.stats.total_low > 0:
            return "LOW", Colors.LOW
        else:
            return "INFORMATIONAL", Colors.INFO


# ============================================================================
# MAIN
# ============================================================================

def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description='Enterprise BurpSuite HTML to DOCX Converter (Professional Edition)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
EXAMPLES:
  %(prog)s report.html
      Convert with all severity levels (default)

  %(prog)s report.html --severity high,medium
      Only include High and Medium severity findings

  %(prog)s report.html --severity high,medium --evidence-only
      Show only evidence/highlighted parts (no full request/response)

  %(prog)s report.html --severity critical,high,medium -o output.docx
      Include Critical/High/Medium (note: BurpSuite maps Critical to High)

  %(prog)s report.html -s high,medium -e --company "MyCompany"
      Shorthand: severity filter + evidence-only mode

SEVERITY LEVELS:
  critical  - Maps to High (BurpSuite highest level)
  high      - High severity vulnerabilities
  medium    - Medium severity vulnerabilities
  low       - Low severity vulnerabilities
  info      - Informational findings

OPTIONS:
  --evidence-only, -e  Show only highlighted evidence portions instead of
                       full HTTP request/response data (smaller report size)
"""
    )

    parser.add_argument('input_file', help='Input BurpSuite HTML file')
    parser.add_argument('-o', '--output', help='Output DOCX file')
    parser.add_argument('--severity', '-s',
                        help='Comma-separated severity levels to include (e.g., high,medium,low). '
                             'Options: critical, high, medium, low, info. Default: all levels')
    parser.add_argument('--evidence-only', '-e', action='store_true',
                        help='Show only evidence/highlighted portions instead of full request/response')
    parser.add_argument('--company', default='Security Assessment Team', help='Company name')
    parser.add_argument('--title', default='Vulnerability Assessment Report', help='Report title')
    parser.add_argument('--target', default='Target Application', help='Target name')
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose output')
    parser.add_argument('--validate', action='store_true',
                        help='Run validation checks after generating report')
    parser.add_argument('--network', '-n',
                        help='Folder containing network pentest Excel files (CVE reports)')

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    if not os.path.exists(args.input_file):
        logger.error(f"File not found: {args.input_file}")
        sys.exit(1)

    output_path = args.output or os.path.splitext(args.input_file)[0] + '_Enterprise_Report.docx'

    logger.info(f"Reading: {args.input_file}")

    with open(args.input_file, 'r', encoding='utf-8', errors='ignore') as f:
        html_content = f.read()

    logger.info(f"Read {len(html_content) / 1024 / 1024:.1f} MB")

    # Validate HTML content
    if not html_content or len(html_content.strip()) == 0:
        logger.error("Error: Input file is empty")
        sys.exit(1)

    if 'BODH0' not in html_content:
        logger.error("Error: Input file does not appear to be a valid BurpSuite HTML report")
        logger.error("Expected to find vulnerability markers (BODH0 class) in the HTML")
        sys.exit(1)

    # Parse severity filter
    severity_filter = None
    if args.severity:
        severity_filter = [s.strip() for s in args.severity.split(',')]
        logger.info(f"Severity filter: {severity_filter}")

    # Parse
    html_parser = BurpHTMLParser(html_content)
    stats, vulns = html_parser.parse()

    # Generate
    generator = DocxReportGenerator(
        statistics=stats,
        vulnerabilities=vulns,
        company_name=args.company,
        report_title=args.title,
        target_name=args.target,
        severity_filter=severity_filter,
        evidence_only=args.evidence_only,
        network_reports_folder=args.network
    )

    generator.generate(output_path)

    total_instances = sum(len(v.instances) for v in vulns)
    filtered_types = len(generator.filtered_vulns)
    filtered_instances = sum(len(v.instances) for v in generator.filtered_vulns)

    print(f"\n{'='*60}")
    print(f"  ENTERPRISE REPORT GENERATED SUCCESSFULLY")
    print(f"{'='*60}")
    print(f"  Output: {output_path}")
    print(f"")
    if severity_filter:
        print(f"  SEVERITY FILTER APPLIED:")
        print(f"  ------------------------")
        print(f"  Filter: {', '.join(severity_filter)}")
        print(f"  Mapped to: {', '.join(sorted(generator.severity_filter))}")
        print(f"")
    print(f"  FINDINGS SUMMARY (ALL):")
    print(f"  ------------------------")
    print(f"  Vulnerability Types : {len(vulns)}")
    print(f"  Total Instances     : {total_instances}")
    print(f"  High Severity       : {stats.total_high}")
    print(f"  Medium Severity     : {stats.total_medium}")
    print(f"  Low Severity        : {stats.total_low}")
    print(f"  Informational       : {stats.total_info}")
    print(f"")
    print(f"  INCLUDED IN REPORT:")
    print(f"  --------------------")
    print(f"  Vulnerability Types : {filtered_types}")
    print(f"  Total Instances     : {filtered_instances}")

    # Network pentest stats
    if generator.network_vulns:
        print(f"")
        print(f"  NETWORK PENTEST (CVE):")
        print(f"  -----------------------")
        print(f"  Domains Scanned     : {len(generator.network_vulns)}")
        total_net = sum(d['total'] for d in generator.network_vulns.values())
        crit_net = sum(d['critical'] for d in generator.network_vulns.values())
        high_net = sum(d['high'] for d in generator.network_vulns.values())
        med_net = sum(d['medium'] for d in generator.network_vulns.values())
        low_net = sum(d['low'] for d in generator.network_vulns.values())
        print(f"  Total CVEs          : {total_net}")
        print(f"  Critical            : {crit_net}")
        print(f"  High                : {high_net}")
        print(f"  Medium              : {med_net}")
        print(f"  Low                 : {low_net}")

    print(f"{'='*60}")

    # Validation check
    if args.validate:
        print(f"\n{'='*60}")
        print(f"  VALIDATION REPORT")
        print(f"{'='*60}")

        validation_passed = True
        issues = []

        # Check 1: Vulnerability count matches BODH0 elements in HTML
        html_bodh0_count = len(html_parser.soup.find_all('span', class_='BODH0'))
        if len(vulns) == html_bodh0_count:
            print(f"  [PASS] Vulnerability types match HTML: {len(vulns)}/{html_bodh0_count}")
        else:
            print(f"  [FAIL] Vulnerability types mismatch: extracted {len(vulns)}, HTML has {html_bodh0_count}")
            validation_passed = False
            issues.append(f"Vulnerability type count mismatch: {len(vulns)} vs {html_bodh0_count}")

        # Check 2: Instance count matches BODH1 elements
        html_bodh1_count = len(html_parser.soup.find_all('span', class_='BODH1'))
        if total_instances == html_bodh1_count:
            print(f"  [PASS] Instance count matches: {total_instances}")
        else:
            print(f"  [WARN] Instance count: {total_instances} vs HTML BODH1: {html_bodh1_count}")
            issues.append(f"Instance count difference: {abs(total_instances - html_bodh1_count)}")

        # Check 3: Severity distribution from overview table
        overview_stats = stats.total_high + stats.total_medium + stats.total_low + stats.total_info
        if overview_stats > 0:
            print(f"  [PASS] Severity stats extracted: {overview_stats} total")
        else:
            print(f"  [WARN] No severity statistics found in overview table")
            issues.append("Overview table stats not found")

        # Check 4: All filtered vulnerabilities have required fields
        missing_fields = 0
        for v in generator.filtered_vulns:
            if not v.title:
                missing_fields += 1
            if not v.severity:
                missing_fields += 1

        if missing_fields == 0:
            print(f"  [PASS] All vulnerabilities have required fields")
        else:
            print(f"  [WARN] {missing_fields} missing required fields")
            issues.append(f"{missing_fields} missing fields")

        # Check 5: Evidence extraction (highlights)
        total_highlights = sum(
            len(inst.request_highlights) + len(inst.response_highlights)
            for v in vulns for inst in v.instances
        )
        print(f"  [INFO] Total evidence highlights extracted: {total_highlights}")

        # Check 6: Unique hosts extracted
        unique_host_count = len(set(
            inst.host for v in vulns for inst in v.instances if inst.host
        ))
        print(f"  [INFO] Unique hosts extracted: {unique_host_count}")

        # Check 7: Request/Response extraction
        rr_div_count = len(html_parser.soup.find_all('div', class_='rr_div'))
        extracted_rr = sum(
            (1 if inst.request else 0) + (1 if inst.response else 0)
            for v in vulns for inst in v.instances
        )
        if rr_div_count > 0:
            extraction_rate = (extracted_rr / rr_div_count * 100) if rr_div_count > 0 else 0
            if extraction_rate >= 90:
                print(f"  [PASS] Request/Response extraction: {extracted_rr}/{rr_div_count} ({extraction_rate:.1f}%)")
            else:
                print(f"  [WARN] Request/Response extraction: {extracted_rr}/{rr_div_count} ({extraction_rate:.1f}%)")
                issues.append(f"Low request/response extraction rate: {extraction_rate:.1f}%")
        else:
            print(f"  [INFO] No request/response divs found in HTML")

        # Check 8: Vulnerabilities without instances (often legitimate - general findings)
        vulns_without_instances = sum(1 for v in vulns if len(v.instances) == 0)
        if vulns_without_instances > 0:
            print(f"  [INFO] General findings (no instances): {vulns_without_instances}")
            print(f"         (These are typically library-level or application-wide findings)")

        # Summary
        print(f"")
        if validation_passed and not issues:
            print(f"  VALIDATION: PASSED - All checks successful")
        elif issues:
            print(f"  VALIDATION: PASSED WITH WARNINGS")
            for issue in issues:
                print(f"    - {issue}")
        else:
            print(f"  VALIDATION: FAILED")

        print(f"{'='*60}")


if __name__ == '__main__':
    main()
